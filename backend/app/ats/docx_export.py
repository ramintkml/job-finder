"""Export structured resume JSON to ATS-friendly DOCX matching Ramin Takmil CV layout."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# Template aligned to C:/Users/Ramin/Downloads/Ramin_Takmil_CV.pdf
FONT_NAME = "Calibri"
COLOR_INK = RGBColor(0x1A, 0x23, 0x30)
COLOR_MUTED = RGBColor(0x5C, 0x67, 0x75)
COLOR_RULE = "C8C4BA"

FA_HEADINGS = {
    "summary": "خلاصه حرفه‌ای",
    "skills": "مهارت‌های اصلی",
    "projects": "پروژه‌های منتخب",
    "additional": "تجربه تکمیلی",
    "experience": "سوابق شغلی",
    "research": "پژوهش",
    "education": "تحصیلات",
    "languages_certs": "زبان‌ها و گواهینامه‌ها",
}


def _set_run_font(
    run,
    *,
    size_pt: float,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    font_name: str = FONT_NAME,
) -> None:
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    # Complex-script (Persian) sizing
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color or COLOR_INK


def _set_paragraph_rtl(paragraph, *, rtl: bool) -> None:
    """Mark paragraph as RTL so Word keeps Persian layout tidy (no manual string reverse)."""
    if not rtl:
        return
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    # Keep bullets from drifting in mixed EN/FA lines
    text_dir = pPr.find(qn("w:textDirection"))
    if text_dir is None:
        text_dir = OxmlElement("w:textDirection")
        pPr.append(text_dir)
    text_dir.set(qn("w:val"), "rl")


def _enable_section_rtl(doc: Document) -> None:
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = sectPr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            sectPr.append(bidi)
        bidi.set(qn("w:val"), "1")


def _set_paragraph_spacing(
    paragraph,
    *,
    before: float = 0,
    after: float = 4,
    line: float = 1.08,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def _add_bottom_border(paragraph) -> None:
    """Subtle rule under section headings (ATS-safe; not a table)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), COLOR_RULE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _tight_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)


def _heading(
    doc: Document,
    text: str,
    *,
    font_name: str = FONT_NAME,
    rtl: bool = False,
) -> None:
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=10, after=4, line=1.05)
    _set_paragraph_rtl(p, rtl=rtl)
    run = p.add_run(text.strip())
    _set_run_font(run, size_pt=11.5, bold=True, font_name=font_name)
    _add_bottom_border(p)


def _body(
    doc: Document,
    text: str,
    *,
    size: float = 10.5,
    after: float = 4,
    font_name: str = FONT_NAME,
    rtl: bool = False,
) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=after, line=1.08)
    _set_paragraph_rtl(p, rtl=rtl)
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, font_name=font_name)


def _bullets(
    doc: Document,
    items: list,
    *,
    font_name: str = FONT_NAME,
    rtl: bool = False,
    level: int = 0,
) -> None:
    indent = 12 + (14 * max(level, 0))
    for item in items:
        if isinstance(item, dict):
            text = str(item.get("text") or item.get("title") or "").strip()
            subs = list(
                item.get("sub_bullets") or item.get("children") or item.get("bullets") or []
            )
            bold = bool(item.get("bold"))
            if text:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=2, after=1, line=1.05)
                _set_paragraph_rtl(p, rtl=rtl)
                p.paragraph_format.left_indent = Pt(indent)
                bullet = "• " if not rtl else "•\u200f "
                run = p.add_run(f"{bullet}{text}")
                _set_run_font(run, size_pt=10.5, bold=bold, font_name=font_name)
            if subs:
                _bullets(doc, subs, font_name=font_name, rtl=rtl, level=level + 1)
            continue
        item = (item or "").strip() if isinstance(item, str) else str(item or "").strip()
        if not item:
            continue
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=2, line=1.05)
        _set_paragraph_rtl(p, rtl=rtl)
        p.paragraph_format.left_indent = Pt(indent)
        bullet = "• " if not rtl else "•\u200f "
        run = p.add_run(f"{bullet}{item}")
        _set_run_font(run, size_pt=10.5, font_name=font_name)


def _role_header(
    doc: Document,
    line: str,
    *,
    font_name: str = FONT_NAME,
    rtl: bool = False,
) -> None:
    line = (line or "").strip()
    if not line:
        return
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=6, after=1, line=1.05)
    _set_paragraph_rtl(p, rtl=rtl)
    run = p.add_run(line)
    _set_run_font(run, size_pt=10.5, bold=True, font_name=font_name)


def _join_bits(*parts: str, sep: str = " | ") -> str:
    return sep.join(p.strip() for p in parts if p and str(p).strip())


def export_resume_docx(
    resume: dict[str, Any],
    path: Path,
    *,
    rtl: bool = False,
    font_name: str | None = None,
    headings: dict[str, str] | None = None,
) -> Path:
    """Export structured resume using the Ramin Takmil CV section order and styling."""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _tight_margins(doc)
    if rtl:
        _enable_section_rtl(doc)

    font = font_name or FONT_NAME
    h = headings or {
        "summary": "Professional Summary",
        "skills": "Core Skills",
        "projects": "Selected Projects",
        "additional": "Additional Experience",
        "experience": "Professional Experience",
        "research": "Research",
        "education": "Education",
        "languages_certs": "Languages & Certifications",
    }
    align = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT

    # --- Header ---
    name = (resume.get("full_name") or "Candidate").strip()
    p = doc.add_paragraph()
    p.alignment = align
    _set_paragraph_spacing(p, before=0, after=0, line=1.0)
    _set_paragraph_rtl(p, rtl=rtl)
    run = p.add_run(name)
    _set_run_font(run, size_pt=18, bold=True, font_name=font)

    title = (resume.get("professional_title") or "").strip()
    if title:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=2, line=1.05)
        _set_paragraph_rtl(p, rtl=rtl)
        run = p.add_run(title)
        _set_run_font(run, size_pt=11.5, bold=True, font_name=font)

    contact = _join_bits(
        str(resume.get("email") or ""),
        str(resume.get("phone") or ""),
    )
    links = _join_bits(
        str(resume.get("linkedin") or ""),
        str(resume.get("github") or ""),
        str(resume.get("portfolio") or ""),
        str(resume.get("location") or ""),
    )
    if contact:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=0 if links else 6, line=1.05)
        _set_paragraph_rtl(p, rtl=rtl)
        run = p.add_run(contact)
        _set_run_font(run, size_pt=9.5, color=COLOR_MUTED, font_name=font)
    if links:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=6, line=1.05)
        _set_paragraph_rtl(p, rtl=rtl)
        run = p.add_run(links)
        _set_run_font(run, size_pt=9.5, color=COLOR_MUTED, font_name=font)

    # --- Professional Summary ---
    summary = (resume.get("summary") or "").strip()
    if summary:
        _heading(doc, h["summary"], font_name=font, rtl=rtl)
        _body(doc, summary, after=6, font_name=font, rtl=rtl)

    # --- Core Skills ---
    skills = resume.get("skills") or {}
    if isinstance(skills, dict) and skills:
        _heading(doc, h["skills"], font_name=font, rtl=rtl)
        for cat, items in skills.items():
            cat = str(cat or "").strip()
            joined = ", ".join(str(x).strip() for x in (items or []) if str(x).strip())
            if not joined:
                continue
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=2, line=1.05)
            _set_paragraph_rtl(p, rtl=rtl)
            if cat:
                run = p.add_run(f"{cat}: ")
                _set_run_font(run, size_pt=10.5, bold=True, font_name=font)
            run = p.add_run(joined)
            _set_run_font(run, size_pt=10.5, font_name=font)
    elif isinstance(skills, list) and skills:
        _heading(doc, h["skills"], font_name=font, rtl=rtl)
        _body(doc, ", ".join(str(x) for x in skills if x), font_name=font, rtl=rtl)

    # --- Selected Projects ---
    projects = resume.get("projects") or []
    if projects:
        _heading(doc, h["projects"], font_name=font, rtl=rtl)
        for proj in projects:
            name_p = (proj.get("name") or "").strip()
            sub = (proj.get("subtitle") or "").strip()
            url = (proj.get("url") or proj.get("link") or "").strip()
            meta = _join_bits(sub, url)
            line = f"{name_p} — {meta}" if meta else name_p
            parent = [{"text": line, "bold": True, "sub_bullets": list(proj.get("bullets") or [])}]
            _bullets(doc, parent, font_name=font, rtl=rtl, level=0)

    # --- Additional Experience ---
    additional = resume.get("additional_experience") or []
    if additional:
        _heading(doc, h["additional"], font_name=font, rtl=rtl)
        for block in additional:
            if isinstance(block, str):
                _body(doc, block, after=3, font_name=font, rtl=rtl)
                continue
            label = (block.get("title") or block.get("label") or "").strip()
            text = (block.get("text") or block.get("summary") or "").strip()
            bullets = list(block.get("bullets") or [])
            if label:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=4, after=1, line=1.05)
                _set_paragraph_rtl(p, rtl=rtl)
                run = p.add_run(label)
                _set_run_font(run, size_pt=10.5, bold=True, font_name=font)
            if text:
                _body(doc, text, after=2, font_name=font, rtl=rtl)
            if bullets:
                _bullets(doc, bullets, font_name=font, rtl=rtl)

    # --- Professional Experience ---
    experience = resume.get("experience") or []
    if experience:
        _heading(doc, h["experience"], font_name=font, rtl=rtl)
        for role in experience:
            title_r = (role.get("title") or "").strip()
            company = (role.get("company") or "").strip()
            location = (role.get("location") or "").strip()
            dates = (role.get("dates") or "").strip()
            left = " - ".join(x for x in [title_r, company] if x)
            line = _join_bits(left, location, dates)
            _role_header(doc, line, font_name=font, rtl=rtl)
            _bullets(doc, list(role.get("bullets") or []), font_name=font, rtl=rtl)

    # --- Research ---
    research = resume.get("research") or []
    if isinstance(research, str) and research.strip():
        _heading(doc, h["research"], font_name=font, rtl=rtl)
        _body(doc, research.strip(), after=4, font_name=font, rtl=rtl)
    elif isinstance(research, list) and research:
        _heading(doc, h["research"], font_name=font, rtl=rtl)
        for item in research:
            if isinstance(item, str):
                _body(doc, item, after=3, font_name=font, rtl=rtl)
            else:
                label = (item.get("title") or "").strip()
                text = (item.get("text") or item.get("summary") or "").strip()
                if label and text:
                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p, before=2, after=2, line=1.05)
                    _set_paragraph_rtl(p, rtl=rtl)
                    run = p.add_run(f"{label}: ")
                    _set_run_font(run, size_pt=10.5, bold=True, font_name=font)
                    run = p.add_run(text)
                    _set_run_font(run, size_pt=10.5, font_name=font)
                elif text or label:
                    _body(doc, text or label, after=3, font_name=font, rtl=rtl)

    # --- Education ---
    education = resume.get("education") or []
    if education:
        _heading(doc, h["education"], font_name=font, rtl=rtl)
        for edu in education:
            degree = (edu.get("degree") or "").strip()
            school = (edu.get("school") or "").strip()
            dates = (edu.get("dates") or "").strip()
            gpa = (edu.get("gpa") or "").strip()
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=2, after=0, line=1.05)
            _set_paragraph_rtl(p, rtl=rtl)
            run = p.add_run(degree or school)
            _set_run_font(run, size_pt=10.5, bold=True, font_name=font)
            detail = _join_bits(school if degree else "", dates, f"GPA {gpa}" if gpa else "")
            if detail:
                p2 = doc.add_paragraph()
                _set_paragraph_spacing(p2, before=0, after=3, line=1.05)
                _set_paragraph_rtl(p2, rtl=rtl)
                run = p2.add_run(detail)
                _set_run_font(run, size_pt=10.5, font_name=font)

    # --- Languages & Certifications ---
    languages = resume.get("languages") or []
    certs = resume.get("certifications") or []
    if languages or certs:
        _heading(doc, h["languages_certs"], font_name=font, rtl=rtl)
        if isinstance(languages, list):
            for lang in languages:
                _body(doc, str(lang).strip(), after=2, font_name=font, rtl=rtl)
        elif isinstance(languages, str) and languages.strip():
            _body(doc, languages.strip(), after=2, font_name=font, rtl=rtl)
        for c in certs:
            _body(doc, str(c).strip(), after=2, font_name=font, rtl=rtl)

    doc.save(str(path))
    return path


def export_resume_docx_fa(resume: dict[str, Any], path: Path) -> Path:
    """Persian RTL DOCX — B Nazanin when bundled/available."""
    from app.ats.fonts import find_bnazanin_font, find_unicode_font, font_family_name

    font_path = find_bnazanin_font() or find_unicode_font(prefer_persian=True)
    family = "B Nazanin"
    if font_path is not None and "nazan" not in font_path.stem.lower():
        family = font_family_name(font_path, fallback="B Nazanin")
    return export_resume_docx(
        resume,
        path,
        rtl=True,
        font_name=family,
        headings=FA_HEADINGS,
    )


def export_markdown_docx(markdown: str, path: Path) -> Path:
    """Parse Codex resume.md into the structured template, then export DOCX."""
    try:
        from app.ats.score import markdown_resume_to_dict

        structured = markdown_resume_to_dict(markdown)
        # Prefer structured template when we got a usable name/summary/experience
        if (structured.get("full_name") or structured.get("summary")) and (
            structured.get("experience") or structured.get("projects") or structured.get("skills")
        ):
            # Enrich from markdown section titles that map to the PDF template
            enriched = _enrich_from_markdown_sections(markdown, structured)
            return export_resume_docx(enriched, path)
    except Exception:
        pass
    return _export_markdown_docx_linear(markdown, path)


def _enrich_from_markdown_sections(markdown: str, base: dict[str, Any]) -> dict[str, Any]:
    """Pull contact + optional template sections from markdown headings."""
    out = dict(base)
    text = (markdown or "").replace("\r\n", "\n")
    lines = text.split("\n")

    # Contact line heuristic near top
    for ln in lines[:12]:
        s = ln.strip()
        if "@" in s and "|" in s:
            parts = [p.strip() for p in s.split("|")]
            for part in parts:
                if "@" in part and not out.get("email"):
                    out["email"] = part
                elif re.search(r"\+?\d[\d\s\-()]{6,}", part) and not out.get("phone"):
                    out["phone"] = part
                elif "linkedin" in part.lower() and not out.get("linkedin"):
                    out["linkedin"] = part
                elif "github" in part.lower() and not out.get("github"):
                    out["github"] = part
                elif re.search(r"(remote|iran|australia|germany|usa|uk)", part, re.I) and not out.get(
                    "location"
                ):
                    out["location"] = part
            break

    # Normalize skills heading label preference
    if out.get("skills") and isinstance(out["skills"], dict):
        # already fine
        pass

    return out


def _export_markdown_docx_linear(markdown: str, path: Path) -> Path:
    """Fallback linear markdown renderer with the same fonts/margins as the template."""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _tight_margins(doc)

    lines = (markdown or "").replace("\r\n", "\n").split("\n")
    i = 0
    first_heading = True
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        i += 1
        if not line:
            continue

        if line.startswith("```"):
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            if i < len(lines):
                i += 1
            continue

        heading_level = 0
        if line.startswith("### "):
            heading_level = 3
            line = line[4:].strip()
        elif line.startswith("## "):
            heading_level = 2
            line = line[3:].strip()
        elif line.startswith("# "):
            heading_level = 1
            line = line[2:].strip()
        elif re.fullmatch(
            r"(Professional Summary|Core Skills|Selected Projects|Additional Experience|"
            r"Professional Experience|Research|Education|Languages & Certifications|"
            r"Summary|Skills|Projects|Certifications)",
            line,
            flags=re.I,
        ):
            heading_level = 2

        if heading_level == 1 and first_heading:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=0, line=1.0)
            run = p.add_run(line)
            _set_run_font(run, size_pt=18, bold=True)
            first_heading = False
            continue
        if heading_level >= 2:
            # Map generic headings to template labels
            mapped = {
                "summary": "Professional Summary",
                "skills": "Core Skills",
                "technical skills": "Core Skills",
                "projects": "Selected Projects",
                "experience": "Professional Experience",
                "work experience": "Professional Experience",
                "certifications": "Languages & Certifications",
            }.get(line.lower(), line)
            _heading(doc, mapped)
            continue
        if heading_level == 3:
            _role_header(doc, line)
            continue

        bullet = False
        indent_level = 0
        stripped = raw.rstrip()
        if stripped.lstrip().startswith(("- ", "* ", "• ")):
            lead = len(stripped) - len(stripped.lstrip())
            indent_level = 1 if lead >= 2 else 0
            bullet = True
            line = stripped.lstrip()[2:].strip()
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in (". ", ") "):
            bullet = True
            line = line[3:].strip() if line[1] == "." else line[2:].strip()

        if bullet:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=2, line=1.05)
            p.paragraph_format.left_indent = Pt(12 + 14 * indent_level)
            _add_inline_runs(p, f"• {line}", size=10.5)
        else:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=3, line=1.05)
            _add_inline_runs(p, line, size=10.5)

    if not doc.paragraphs:
        p = doc.add_paragraph()
        run = p.add_run("(empty resume)")
        _set_run_font(run, size_pt=10.5)

    doc.save(str(path))
    return path


def _add_inline_runs(paragraph, text: str, *, size: float = 10.5) -> None:
    """Parse light **bold** / *italic* markers into runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text or ""):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            _set_run_font(run, size_pt=size)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, size_pt=size, bold=True)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, size_pt=size, italic=True)
        else:
            run = paragraph.add_run(token.strip("`"))
            _set_run_font(run, size_pt=size)
        pos = match.end()
    if pos < len(text or ""):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, size_pt=size)
    if not (text or "").strip() and not paragraph.runs:
        run = paragraph.add_run("")
        _set_run_font(run, size_pt=size)
