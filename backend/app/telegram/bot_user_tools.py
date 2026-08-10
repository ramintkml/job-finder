"""Per-user AI resume + ATS tools (Telegram), gated by subscription add-ons."""

from __future__ import annotations

import html
import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any

from app.billing.service import get_user_by_telegram_id, user_entitlements
from app.config import DATA_DIR
from app.database import SessionLocal
from app.telegram import keyboards as kb

logger = logging.getLogger(__name__)

USER_CV_DIR = DATA_DIR / "user_cvs"
USER_OUT_DIR = DATA_DIR / "user_outputs"

# telegram_user_id -> pending tool state
_pending: dict[int, dict[str, Any]] = {}
# telegram_user_id -> last generated package (for ATS improve button)
_last_runs: dict[int, dict[str, Any]] = {}

AI_RESUME_LABELS = {
    "📄 رزومه با AI",
    "رزومه با AI",
    "/ai_resume",
}

ATS_CHECK_LABELS = {
    "📊 بررسی ATS",
    "بررسی ATS",
    "/ats",
}


def application_actions_markup(app, *, improve_btn: dict | None = None) -> dict:
    """Inline buttons for Codex/AI resumes: Improve, FA translate, PDF(s)."""
    app_id = int(app.id)
    has_fa = bool(
        getattr(app, "resume_fa_docx_path", None)
        or getattr(app, "resume_fa_json", None)
        or getattr(app, "resume_fa_pdf_path", None)
    )
    if improve_btn is None:
        improve_btn = kb._btn("🔄 Improve (evaluation)", f"ut:improve_codex:{app_id}")
    rows = [
        [
            improve_btn,
            kb._btn("🇮🇷 ترجمه فارسی", f"ut:fa:{app_id}"),
        ]
    ]
    pdf_row = [kb._btn("📄 PDF English", f"ut:pdf_en:{app_id}")]
    if has_fa:
        pdf_row.append(kb._btn("📄 PDF فارسی", f"ut:pdf_fa:{app_id}"))
    rows.append(pdf_row)
    return kb._markup(rows)


def _load_app_for_user(db, app_id: int, telegram_user_id: int):
    from app.database import BotApplication

    app = db.get(BotApplication, app_id)
    if not app or int(app.telegram_user_id) != int(telegram_user_id):
        return None
    return app


def _parse_resume_struct(app) -> dict[str, Any]:
    import json

    if app.resume_json:
        try:
            data = json.loads(app.resume_json)
            if isinstance(data, dict) and (data.get("full_name") or data.get("summary")):
                return data
        except Exception:
            pass
    md = (app.resume_md or "").strip()
    if not md:
        raise ValueError("No resume stored for this application")
    from app.ats.score import markdown_resume_to_dict

    structured = markdown_resume_to_dict(md)
    try:
        from app.ats.docx_export import _enrich_from_markdown_sections

        structured = _enrich_from_markdown_sections(md, structured)
    except Exception:
        pass
    return structured


def pending_for(telegram_user_id: int) -> dict[str, Any] | None:
    return _pending.get(int(telegram_user_id))


def clear_pending(telegram_user_id: int) -> None:
    _pending.pop(int(telegram_user_id), None)


def _esc(v: Any) -> str:
    return html.escape(str(v if v is not None else "—"))


def _user_cv_path(telegram_user_id: int) -> Path:
    return USER_CV_DIR / str(telegram_user_id) / "base_cv.txt"


def load_user_cv_text(telegram_user_id: int) -> str | None:
    path = _user_cv_path(telegram_user_id)
    if path.is_file():
        text = path.read_text(encoding="utf-8").strip()
        return text or None
    return None


def save_user_cv_text(telegram_user_id: int, text: str) -> Path:
    path = _user_cv_path(telegram_user_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip(), encoding="utf-8")
    return path


def extract_text_from_docx_bytes(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def parse_linkedin_job_id(text: str) -> str | None:
    raw = (text or "").strip()
    match = re.search(r"/jobs/view/(?:[^/?#]+-)?(\d+)", raw)
    if match:
        return match.group(1)
    match = re.search(r"linkedin\.com/jobs/view/(\d+)", raw, re.I)
    if match:
        return match.group(1)
    if raw.isdigit() and len(raw) >= 8:
        return raw
    return None


def _entitlements_for(telegram_user_id: int):
    db = SessionLocal()
    try:
        user = get_user_by_telegram_id(db, telegram_user_id)
        if not user:
            return None, None
        return user, user_entitlements(db, user.id)
    finally:
        db.close()


COMBINED_LABELS = {
    "📄📊 رزومه + ATS",
    "رزومه + ATS",
    "رزومه و ATS",
    "/resume_ats",
    "/combined",
}

APPLY_LABELS = {
    "📋 ارسال آگهی",
    "ارسال آگهی",
    "📝 درخواست شغل",
    "درخواست شغل",
    "/apply",
    "apply",
}

PC_WORKER_LABELS = {
    "🖥 PC Worker",
    "PC Worker",
    "وضعیت PC Worker",
    "/worker",
    "worker",
}


def pc_worker_status_message() -> dict:
    """Reply-keyboard action: show whether the PC worker is online."""
    from app.database import SessionLocal
    from app.worker.queue import worker_status

    db = SessionLocal()
    try:
        ws = worker_status(db)
    finally:
        db.close()
    online = bool(ws.get("worker_online"))
    icon = _bool_icon(online)
    lines = [
        "<b>وضعیت PC Worker</b>",
        "",
        f"PC worker: {icon} {'آنلاین' if online else 'آفلاین'}",
    ]
    if not online:
        lines.extend(
            [
                "",
                "برای آنلاین شدن:",
                "۱) تونل SSH به VPS",
                "۲) اجرای <code>launch.bat</code> روی PC",
            ]
        )
    name = (ws.get("worker_name") or ws.get("worker_id") or "").strip()
    if name:
        lines.append(f"شناسه: <code>{html.escape(name[:64])}</code>")
    return {"message": "\n".join(lines)}


def _bool_icon(v: bool) -> str:
    return "✅" if v else "❌"


def _is_admin_user(telegram_user_id: int) -> bool:
    from app.config import settings

    return int(telegram_user_id) in settings.admin_telegram_ids


def _resolve_base_cv(telegram_user_id: int, *, pending: dict | None = None) -> str:
    """Prefer user DOCX CV; admin may fall back to operator LinkedIn/base CV."""
    if pending:
        text = str(pending.get("cv_text") or "").strip()
        if len(text) >= 40:
            return text
    saved = load_user_cv_text(telegram_user_id)
    if saved and len(saved) >= 40:
        return saved
    if _is_admin_user(telegram_user_id):
        try:
            from app.ats.tailor import load_base_cv_text

            return load_base_cv_text().strip()
        except Exception:
            logger.exception("Admin base CV fallback failed")
        try:
            from app.database import SessionLocal
            from app.linkedin.settings import load_linkedin_settings

            db = SessionLocal()
            try:
                cfg = load_linkedin_settings(db)
            finally:
                db.close()
            if (cfg.cv_text or "").strip():
                return cfg.cv_text.strip()
        except Exception:
            logger.exception("Admin LinkedIn CV fallback failed")
    return ""


def start_apply(telegram_user_id: int, *, chat_id: int | None = None) -> dict:
    """Queue JD for local Codex / job-search-copilot (preferred), with VPS Groq fallback message."""
    user, ent = _entitlements_for(telegram_user_id)
    is_admin = _is_admin_user(telegram_user_id)
    allowed = is_admin or (
        user is not None
        and ent is not None
        and ent.has_active_plan
        and bool(ent.include_ai)
    )
    if not allowed:
        return {
            "message": (
                "❌ برای <b>/apply</b> به اشتراک با افزونه <b>هوش مصنوعی</b> نیاز دارید.\n"
                "از <b>💎 پلن‌ها</b> یا <b>⬆️ ارتقای پلن</b> استفاده کنید."
            ),
            "reply_markup": kb._markup(
                [[kb._btn("💎 پلن‌ها", "pl:home")], [kb._btn("⬆️ ارتقای پلن", "pl:up")]]
            ),
        }

    cancel_kb = kb._markup([[kb._btn("لغو", "ut:cancel")]])
    base = _resolve_base_cv(telegram_user_id)
    # Codex uses Job Search profile/; local CV still helpful context for fallback
    _pending[int(telegram_user_id)] = {
        "flow": "apply_codex",
        "step": "job",
        "cv_text": base if len(base) >= 40 else "",
        "chat_id": int(chat_id or telegram_user_id),
    }
    from app.database import SessionLocal
    from app.worker.queue import worker_status

    db = SessionLocal()
    try:
        ws = worker_status(db)
    finally:
        db.close()
    worker_line = (
        f"PC worker: {_bool_icon(bool(ws.get('worker_online')))} "
        + ("آنلاین" if ws.get("worker_online") else "آفلاین")
    )
    return {
        "message": (
            "لطفا آگهی شغلی یا لینک linkedin رو ارسال کنید.\n\n"
            f"{worker_line}"
        ),
        "reply_markup": cancel_kb,
    }


def tool_flags_for_user(telegram_user_id: int | None) -> tuple[bool, bool]:
    """Return (include_ai, include_ats) for active plan; else (False, False)."""
    if telegram_user_id is None:
        return False, False
    if _is_admin_user(int(telegram_user_id)):
        return True, True
    _user, ent = _entitlements_for(int(telegram_user_id))
    if ent is None or not ent.has_active_plan:
        return False, False
    return bool(ent.include_ai), bool(ent.include_ats)



def start_ai_resume(telegram_user_id: int) -> dict:
    user, ent = _entitlements_for(telegram_user_id)
    if user is None or ent is None or not ent.has_active_plan:
        return {
            "message": (
                "❌ اشتراک فعالی ندارید.\n"
                "از <b>💎 پلن‌ها</b> خرید کنید، سپس افزونه <b>هوش مصنوعی</b> را فعال کنید."
            ),
        }
    if ent.include_ai and ent.include_ats:
        return start_combined(telegram_user_id)
    if not ent.include_ai:
        return {
            "message": (
                "❌ افزونه <b>هوش مصنوعی</b> در پلن شما فعال نیست.\n"
                "از <b>⬆️ ارتقای پلن</b> هوش مصنوعی را اضافه کنید."
            ),
            "reply_markup": kb._markup(
                [[kb._btn("⬆️ ارتقای پلن", "pl:up")], [kb._btn("💎 پلن‌ها", "pl:home")]]
            ),
        }

    has_cv = bool(load_user_cv_text(telegram_user_id))
    cancel_kb = kb._markup([[kb._btn("لغو", "ut:cancel")]])
    if not has_cv:
        _pending[int(telegram_user_id)] = {"flow": "ai", "step": "cv"}
        return {
            "message": (
                "<b>📄 رزومه با AI</b>\n\n"
                "اول <b>رزومه پایه</b> را به‌صورت فایل <b>DOCX</b> بفرستید "
                "(یک‌بار ذخیره می‌شود).\n"
                "بعد توضیح شغل یا لینک لینکدین را می‌فرستید."
            ),
            "reply_markup": cancel_kb,
        }

    _pending[int(telegram_user_id)] = {"flow": "ai", "step": "job"}
    return {
        "message": (
            "<b>📄 رزومه با AI</b>\n\n"
            "توضیح شغل را بفرستید، یا <b>لینک آگهی لینکدین</b>.\n"
            "رزومه متناسب با شغل برایتان ساخته می‌شود.\n\n"
            "برای تعویض رزومه پایه، یک DOCX جدید بفرستید."
        ),
        "reply_markup": cancel_kb,
    }


def start_ats_check(telegram_user_id: int) -> dict:
    user, ent = _entitlements_for(telegram_user_id)
    if user is None or ent is None or not ent.has_active_plan:
        return {
            "message": (
                "❌ اشتراک فعالی ندارید.\n"
                "از <b>💎 پلن‌ها</b> خرید کنید، سپس افزونه <b>ATS</b> را فعال کنید."
            ),
        }
    if ent.include_ai and ent.include_ats:
        return start_combined(telegram_user_id)
    if not ent.include_ats:
        return {
            "message": (
                "❌ افزونه <b>ATS</b> در پلن شما فعال نیست.\n"
                "از <b>⬆️ ارتقای پلن</b> ATS را اضافه کنید."
            ),
            "reply_markup": kb._markup(
                [[kb._btn("⬆️ ارتقای پلن", "pl:up")], [kb._btn("💎 پلن‌ها", "pl:home")]]
            ),
        }

    _pending[int(telegram_user_id)] = {"flow": "ats", "step": "cv"}
    return {
        "message": (
            "<b>📊 بررسی ATS</b>\n\n"
            "۱. فایل رزومه را به‌صورت <b>DOCX</b> بفرستید.\n"
            "۲. سپس توضیح شغل (یا لینک لینکدین) را بفرستید.\n\n"
            "امتیاز ATS و بهبودهای قابل اعمال برایتان ارسال می‌شود.\n"
            "رزومه جدیدی ساخته نمی‌شود (ساخت رزومه فقط با افزونه AI)."
        ),
        "reply_markup": kb._markup([[kb._btn("لغو", "ut:cancel")]]),
    }


def start_combined(telegram_user_id: int) -> dict:
    """Single flow when user has both AI and ATS: DOCX + JD/link → resume + ATS report."""
    user, ent = _entitlements_for(telegram_user_id)
    if user is None or ent is None or not ent.has_active_plan:
        return {
            "message": (
                "❌ اشتراک فعالی ندارید.\n"
                "از <b>💎 پلن‌ها</b> پلن با <b>هوش مصنوعی</b> و <b>ATS</b> بخرید."
            ),
        }
    if not (ent.include_ai and ent.include_ats):
        missing = []
        if not ent.include_ai:
            missing.append("هوش مصنوعی")
        if not ent.include_ats:
            missing.append("ATS")
        return {
            "message": (
                "❌ برای حالت ترکیبی به هر دو افزونه نیاز است.\n"
                f"کمبود: <b>{' + '.join(missing)}</b>\n"
                "از <b>⬆️ ارتقای پلن</b> اضافه کنید."
            ),
            "reply_markup": kb._markup(
                [[kb._btn("⬆️ ارتقای پلن", "pl:up")], [kb._btn("💎 پلن‌ها", "pl:home")]]
            ),
        }

    has_cv = bool(load_user_cv_text(telegram_user_id))
    cancel_kb = kb._markup([[kb._btn("لغو", "ut:cancel")]])
    if not has_cv:
        _pending[int(telegram_user_id)] = {"flow": "combined", "step": "cv"}
        return {
            "message": (
                "<b>📄📊 رزومه + ATS</b>\n\n"
                "چون هر دو افزونه فعال است، یک‌جا انجام می‌شود:\n"
                "۱. رزومه را به‌صورت <b>DOCX</b> بفرستید\n"
                "۲. توضیح شغل یا لینک لینکدین را بفرستید\n\n"
                "خروجی: رزومه متناسب با شغل + امتیاز و راهنمای ATS"
            ),
            "reply_markup": cancel_kb,
        }

    _pending[int(telegram_user_id)] = {"flow": "combined", "step": "job"}
    return {
        "message": (
            "<b>📄📊 رزومه + ATS</b>\n\n"
            "توضیح شغل یا <b>لینک لینکدین</b> را بفرستید.\n"
            "رزومه AI + گزارش ATS با هم ساخته می‌شود.\n\n"
            "برای تعویض رزومه پایه، یک DOCX جدید بفرستید."
        ),
        "reply_markup": cancel_kb,
    }


async def handle_tool_callback(data: str, *, telegram_user_id: int) -> dict:
    """Handle ut:* callbacks (user tools)."""
    if data == "ut:cancel":
        if pending_for(telegram_user_id) is None:
            return {"toast": "چیزی برای لغو نیست"}
        clear_pending(telegram_user_id)
        return {
            "toast": "لغو شد",
            "edit_text": "لغو شد.",
            "reply_markup": kb.cleared_keyboard(),
        }
    if data.startswith("ut:improve_codex:"):
        return await _handle_improve_codex(data, telegram_user_id=telegram_user_id)
    if data.startswith("ut:fa:"):
        return await _handle_translate_fa(data, telegram_user_id=telegram_user_id)
    if data.startswith("ut:pdf_en:"):
        return await _handle_pdf(data, telegram_user_id=telegram_user_id, lang="en")
    if data.startswith("ut:pdf_fa:"):
        return await _handle_pdf(data, telegram_user_id=telegram_user_id, lang="fa")
    if data == "ut:improve":
        last = _last_runs.get(int(telegram_user_id))
        if not last:
            return {
                "toast": "بسته قبلی پیدا نشد",
                "alert": True,
                "message": (
                    "❌ هنوز خروجی رزومه‌ای برای بهبود نیست.\n"
                    "اول /apply یا رزومه AI را اجرا کنید."
                ),
            }
        tips = last.get("tips")
        scoring = last.get("scoring")
        if not tips and not scoring:
            return {
                "toast": "نکته ATS نیست",
                "alert": True,
                "message": "❌ نکته‌ای از ATS برای بهبود ذخیره نشده است.",
            }
        return {
            "toast": "در حال بهبود…",
            "message": (
                "⏳ در حال بازتولید رزومه با تمرکز روی "
                "<b>کلیدواژه‌ها و بخش‌های جاافتاده ATS</b>…"
            ),
            "then_generate": {
                "telegram_user_id": telegram_user_id,
                "base_cv": last.get("base_cv") or "",
                "title": last.get("title") or "",
                "company": last.get("company") or "",
                "job_url": last.get("job_url") or "",
                "description": last.get("description") or "",
                "mode": last.get("mode") or "apply",
                "prior_tips": tips,
                "prior_scoring": scoring,
                "improve_pass": True,
                "previous_ats_score": int((scoring or {}).get("total_score") or 0),
            },
        }
    return {"toast": "نامشخص"}


async def _handle_improve_codex(data: str, *, telegram_user_id: int) -> dict:
    from app.ats.tips import build_improvement_guide, format_guidance_for_tailor
    from app.database import BotApplication
    from app.worker.queue import JOB_CODEX_APPLY, enqueue_work, worker_status

    try:
        app_id = int(data.rsplit(":", 1)[-1])
    except ValueError:
        return {"toast": "شناسه نامعتبر", "alert": True}

    db = SessionLocal()
    try:
        app = db.get(BotApplication, app_id)
        if not app or int(app.telegram_user_id) != int(telegram_user_id):
            return {
                "toast": "پیدا نشد",
                "alert": True,
                "message": "❌ این درخواست رزومه پیدا نشد یا متعلق به شما نیست.",
            }

        evaluation_md = (app.evaluation_md or "").strip()
        if not evaluation_md and app.evaluation_path:
            try:
                from pathlib import Path

                p = Path(app.evaluation_path)
                if p.is_file():
                    evaluation_md = p.read_text(encoding="utf-8")
                    app.evaluation_md = evaluation_md[:200000]
            except Exception:
                logger.exception("Failed reading evaluation_path for app %s", app_id)

        description = (app.description or "").strip()
        if len(description) < 40:
            return {
                "toast": "توضیح شغل نیست",
                "alert": True,
                "message": "❌ توضیح شغل برای بهبود ذخیره نشده است. یک بار دیگر آگهی را بفرستید.",
            }
        if not evaluation_md and not (app.resume_md or "").strip():
            return {
                "toast": "evaluation نیست",
                "alert": True,
                "message": "❌ evaluation.md / resume برای این شغل ذخیره نشده است.",
            }

        ats_guidance = ""
        scoring = None
        if app.ats_scores_json:
            try:
                import json as _json

                scoring = _json.loads(app.ats_scores_json)
            except Exception:
                scoring = None
        if scoring:
            tips = build_improvement_guide(scoring)
            ats_guidance = format_guidance_for_tailor(tips, scoring)

        app.status = "improving"
        db.commit()

        ws = worker_status(db)
        work = enqueue_work(
            db,
            JOB_CODEX_APPLY,
            int(telegram_user_id),
            extra={
                "title": app.title or "Target role",
                "company": app.company or "",
                "job_url": app.job_url or "",
                "description": description,
                "chat_id": int(telegram_user_id),
                "telegram_user_id": int(telegram_user_id),
                "improve": True,
                "application_id": app.id,
                "evaluation_md": evaluation_md,
                "previous_resume_md": app.resume_md or "",
                "previous_output_dir": app.output_dir or "",
                "ats_guidance": ats_guidance[:12000],
                "previous_ats_score": int(app.ats_score or 0),
            },
            dedupe=False,
        )
    finally:
        db.close()

    online = bool(ws.get("worker_online"))
    msg = (
        "⏳ درخواست <b>Improve</b> به صف Codex اضافه شد.\n"
        f"شناسه صف: <code>#{work.id}</code>\n"
        "رزومه دوباره ساخته می‌شود با تمرکز روی "
        "<b>evaluation.md</b> و <b>شکاف‌های ATS</b>.\n\n"
    )
    if online:
        msg += "PC Worker آنلاین است — چند دقیقه صبر کنید."
    else:
        msg += (
            "⚠️ PC Worker آفلاین است.\n"
            "روی PC فایل <code>launch.bat</code> را اجرا کنید."
        )
    return {"toast": "Improve صف شد", "message": msg}


async def _handle_translate_fa(data: str, *, telegram_user_id: int) -> dict:
    import json

    from app.ats.docx_export import export_resume_docx_fa
    from app.ats.pdf_export import export_resume_pdf_fa
    from app.ats.translate_fa import translate_resume_to_persian
    from app.config import ATS_DIR

    try:
        app_id = int(data.rsplit(":", 1)[-1])
    except ValueError:
        return {"toast": "شناسه نامعتبر", "alert": True}

    db = SessionLocal()
    try:
        app = _load_app_for_user(db, app_id, telegram_user_id)
        if not app:
            return {
                "toast": "پیدا نشد",
                "alert": True,
                "message": "❌ این رزومه پیدا نشد یا متعلق به شما نیست.",
            }
        try:
            en = _parse_resume_struct(app)
        except Exception as exc:
            return {
                "toast": "رزومه نیست",
                "alert": True,
                "message": f"❌ رزومه انگلیسی برای ترجمه موجود نیست.\n<code>{html.escape(str(exc)[:120])}</code>",
            }

        try:
            fa = await translate_resume_to_persian(en)
        except Exception as exc:
            logger.exception("FA translate failed for app %s", app_id)
            return {
                "toast": "خطای ترجمه",
                "alert": True,
                "message": f"❌ ترجمه فارسی ناموفق بود.\n<code>{html.escape(str(exc)[:160])}</code>",
            }

        out_dir = Path(ATS_DIR) / "applications" / str(app.id)
        out_dir.mkdir(parents=True, exist_ok=True)
        from app.ats.naming import resume_filename, sanitize_job_title

        job_title = sanitize_job_title(app.title or "Role", max_len=60)
        fa_docx = out_dir / resume_filename(lang="fa", job_title=job_title, ext="docx")
        fa_pdf = out_dir / resume_filename(lang="fa", job_title=job_title, ext="pdf")
        # Keep legacy aliases for older buttons/paths
        legacy_fa_docx = out_dir / "resume_fa.docx"
        legacy_fa_pdf = out_dir / "resume_fa.pdf"
        export_resume_docx_fa(fa, fa_docx)
        if fa_docx.is_file():
            legacy_fa_docx.write_bytes(fa_docx.read_bytes())
        pdf_ok = False
        try:
            export_resume_pdf_fa(fa, fa_pdf)
            pdf_ok = fa_pdf.is_file()
            if pdf_ok:
                legacy_fa_pdf.write_bytes(fa_pdf.read_bytes())
        except Exception:
            logger.exception("FA PDF export failed for app %s", app_id)

        app.resume_fa_json = json.dumps(fa, ensure_ascii=False, default=str)[:400000]
        if fa_docx.is_file():
            app.resume_fa_docx_path = str(fa_docx)
        if pdf_ok:
            app.resume_fa_pdf_path = str(fa_pdf)
        # Ensure EN PDF exists too (for dual PDF buttons)
        if not (app.resume_pdf_path and Path(app.resume_pdf_path).is_file()):
            try:
                from app.ats.pdf_export import export_resume_pdf

                en_pdf = out_dir / resume_filename(lang="en", job_title=job_title, ext="pdf")
                export_resume_pdf(en, en_pdf)
                if en_pdf.is_file():
                    app.resume_pdf_path = str(en_pdf)
                    (out_dir / "resume.pdf").write_bytes(en_pdf.read_bytes())
            except Exception:
                logger.exception("EN PDF backfill failed for app %s", app_id)
        if not app.resume_json:
            app.resume_json = json.dumps(en, ensure_ascii=False, default=str)[:400000]

        # Sync FA artifacts into the PC Job Search applications folder via worker
        sync_note = ""
        pc_out = (app.output_dir or "").strip()
        if pc_out and fa_docx.is_file():
            try:
                import base64

                from app.worker.queue import JOB_SAVE_FILES, enqueue_work, worker_status

                files_b64: dict[str, str] = {
                    fa_docx.name: base64.b64encode(fa_docx.read_bytes()).decode("ascii"),
                    # legacy names too
                    "resume_fa.docx": base64.b64encode(fa_docx.read_bytes()).decode("ascii"),
                }
                if pdf_ok:
                    files_b64[fa_pdf.name] = base64.b64encode(fa_pdf.read_bytes()).decode(
                        "ascii"
                    )
                    files_b64["resume_fa.pdf"] = files_b64[fa_pdf.name]
                html_side = fa_pdf.with_suffix(".html")
                if html_side.is_file():
                    files_b64[html_side.name] = base64.b64encode(
                        html_side.read_bytes()
                    ).decode("ascii")
                files_b64["resume_fa.json"] = base64.b64encode(
                    app.resume_fa_json.encode("utf-8")
                ).decode("ascii")
                work = enqueue_work(
                    db,
                    JOB_SAVE_FILES,
                    int(telegram_user_id),
                    extra={
                        "output_dir": pc_out,
                        "files": files_b64,
                        "application_id": app.id,
                    },
                    dedupe=False,
                )
                ws = worker_status(db)
                if ws.get("worker_online"):
                    sync_note = (
                        f"\n📁 در حال ذخیره در پوشه PC "
                        f"(صف <code>#{work.id}</code>)…"
                    )
                else:
                    sync_note = (
                        "\n⚠️ برای ذخیره در پوشه Job Search روی PC، "
                        "<code>launch.bat</code> را اجرا کنید "
                        f"(صف <code>#{work.id}</code>)."
                    )
            except Exception as exc:
                logger.exception("Failed enqueueing FA file sync for app %s", app_id)
                sync_note = f"\n⚠️ ذخیره PC صف نشد: <code>{html.escape(str(exc)[:120])}</code>"
        elif not pc_out:
            sync_note = (
                "\nℹ️ مسیر پوشه PC برای این شغل ذخیره نشده "
                "(فقط خروجی‌های Codex /apply همگام می‌شوند)."
            )

        db.commit()
        db.refresh(app)
        markup = application_actions_markup(app)
        docs: list[dict[str, str]] = []
        if fa_docx.is_file():
            docs.append(
                {
                    "path": str(fa_docx),
                    "caption": f"{fa_docx.name} (RTL · B Nazanin)",
                }
            )
        if pdf_ok:
            docs.append(
                {
                    "path": str(fa_pdf),
                    "caption": f"{fa_pdf.name} (HTML→PDF · B Nazanin)",
                }
            )
        return {
            "toast": "ترجمه آماده",
            "message": (
                "✅ <b>ترجمه فارسی</b> آماده شد (DOCX راست‌چین · فونت B Nazanin).\n"
                "PDF از HTML ساخته می‌شود تا جملات و بولت‌ها درست بمانند.\n"
                "حالا دکمه <b>PDF فارسی</b> هم فعال است."
                f"{sync_note}"
            ),
            "reply_markup": markup,
            "send_documents": docs,
            "skip_improve_hint": True,
        }
    finally:
        db.close()


async def _handle_pdf(data: str, *, telegram_user_id: int, lang: str) -> dict:
    import json

    from app.config import ATS_DIR

    try:
        app_id = int(data.rsplit(":", 1)[-1])
    except ValueError:
        return {"toast": "شناسه نامعتبر", "alert": True}

    db = SessionLocal()
    try:
        app = _load_app_for_user(db, app_id, telegram_user_id)
        if not app:
            return {
                "toast": "پیدا نشد",
                "alert": True,
                "message": "❌ این رزومه پیدا نشد یا متعلق به شما نیست.",
            }

        out_dir = Path(ATS_DIR) / "applications" / str(app.id)
        out_dir.mkdir(parents=True, exist_ok=True)
        markup = application_actions_markup(app)
        from app.ats.naming import resume_filename, sanitize_job_title

        job_title = sanitize_job_title(app.title or "Role", max_len=60)

        if lang == "fa":
            named = out_dir / resume_filename(lang="fa", job_title=job_title, ext="pdf")
            path = Path(app.resume_fa_pdf_path) if app.resume_fa_pdf_path else named
            if not path.is_file():
                path = named
                fa = None
                if app.resume_fa_json:
                    try:
                        fa = json.loads(app.resume_fa_json)
                    except Exception:
                        fa = None
                if not isinstance(fa, dict):
                    return {
                        "toast": "فارسی نیست",
                        "alert": True,
                        "message": "❌ هنوز ترجمه فارسی ندارید. اول دکمه <b>🇮🇷 ترجمه فارسی</b> را بزنید.",
                    }
                try:
                    from app.ats.pdf_export import export_resume_pdf_fa

                    export_resume_pdf_fa(fa, path)
                    app.resume_fa_pdf_path = str(path)
                    db.commit()
                except Exception as exc:
                    logger.exception("FA PDF generate failed")
                    return {
                        "toast": "خطای PDF",
                        "alert": True,
                        "message": f"❌ ساخت PDF فارسی ناموفق بود.\n<code>{html.escape(str(exc)[:160])}</code>",
                    }
            # Ensure Telegram shows the canonical filename
            if path.name != named.name and path.is_file():
                named.write_bytes(path.read_bytes())
                path = named
                app.resume_fa_pdf_path = str(path)
                db.commit()
            return {
                "toast": "PDF فارسی",
                "send_documents": [
                    {"path": str(path), "caption": path.name}
                ],
                "reply_markup": markup,
                "skip_improve_hint": True,
            }

        # English
        named = out_dir / resume_filename(lang="en", job_title=job_title, ext="pdf")
        path = Path(app.resume_pdf_path) if app.resume_pdf_path else named
        if not path.is_file():
            path = named
            try:
                en = _parse_resume_struct(app)
                from app.ats.pdf_export import export_resume_pdf

                export_resume_pdf(en, path)
                app.resume_pdf_path = str(path)
                if not app.resume_json:
                    app.resume_json = json.dumps(en, ensure_ascii=False, default=str)[:400000]
                db.commit()
            except Exception as exc:
                logger.exception("EN PDF generate failed")
                return {
                    "toast": "خطای PDF",
                    "alert": True,
                    "message": f"❌ ساخت PDF انگلیسی ناموفق بود.\n<code>{html.escape(str(exc)[:160])}</code>",
                }
        if path.name != named.name and path.is_file():
            named.write_bytes(path.read_bytes())
            path = named
            app.resume_pdf_path = str(path)
            db.commit()
        return {
            "toast": "PDF English",
            "send_documents": [
                {"path": str(path), "caption": path.name}
            ],
            "reply_markup": markup,
            "skip_improve_hint": True,
        }
    finally:
        db.close()


async def handle_cancel(telegram_user_id: int) -> dict | None:
    if pending_for(telegram_user_id) is None:
        return None
    clear_pending(telegram_user_id)
    return {"handled": True, "message": "لغو شد."}


async def handle_document(
    *,
    telegram_user_id: int,
    filename: str,
    content: bytes,
) -> dict:
    pending = pending_for(telegram_user_id)
    if not pending:
        return {"handled": False}

    name = (filename or "").lower()
    if not name.endswith(".docx"):
        return {
            "handled": True,
            "message": "❌ لطفاً فقط فایل <b>DOCX</b> بفرستید (Word).",
        }
    try:
        text = extract_text_from_docx_bytes(content)
    except Exception as exc:
        logger.exception("DOCX parse failed")
        return {"handled": True, "message": f"❌ خواندن DOCX ممکن نشد: {_esc(exc)}"}
    if len(text) < 40:
        return {"handled": True, "message": "❌ متن استخراج‌شده از DOCX خیلی کوتاه است."}

    flow = pending.get("flow")
    cancel_kb = kb._markup([[kb._btn("لغو", "ut:cancel")]])
    if flow in ("ai", "combined", "apply"):
        save_user_cv_text(telegram_user_id, text)
        pending["cv_text"] = text
        pending["step"] = "job"
        _pending[int(telegram_user_id)] = pending
        label = "رزومه پایه" if flow in ("ai", "apply") else "رزومه"
        return {
            "handled": True,
            "message": (
                f"✅ {label} ذخیره شد ({len(text):,} کاراکتر).\n\n"
                "حالا <b>توضیح شغل</b> یا <b>لینک لینکدین</b> را بفرستید."
            ),
            "reply_markup": cancel_kb,
        }

    if flow == "ats" and pending.get("step") == "cv":
        pending["cv_text"] = text
        pending["step"] = "job"
        _pending[int(telegram_user_id)] = pending
        return {
            "handled": True,
            "message": (
                f"✅ رزومه دریافت شد ({len(text):,} کاراکتر).\n\n"
                "حالا <b>توضیح شغل</b> یا <b>لینک لینکدین</b> را بفرستید."
            ),
            "reply_markup": cancel_kb,
        }

    return {"handled": True, "message": "در این مرحله متن شغل را بفرستید (نه فایل)."}


async def handle_text(*, telegram_user_id: int, text: str) -> dict:
    pending = pending_for(telegram_user_id)
    if not pending:
        return {"handled": False}

    raw = (text or "").strip()
    if raw.lower() in ("cancel", "/cancel", "لغو", "انصراف"):
        clear_pending(telegram_user_id)
        return {"handled": True, "message": "لغو شد."}

    if pending.get("step") == "cv":
        return {
            "handled": True,
            "message": "لطفاً اول فایل <b>DOCX</b> رزومه را بفرستید (یا <code>لغو</code>).",
        }

    if pending.get("step") != "job":
        clear_pending(telegram_user_id)
        return {"handled": True, "message": "وضعیت نامشخص — دوباره از منو شروع کنید."}

    job_title, job_company, job_url, description = await _resolve_job_input(raw)
    if len(description) < 40:
        return {
            "handled": True,
            "message": (
                "❌ توضیح شغل خیلی کوتاه است یا لینک لینکدین خوانده نشد.\n"
                "متن کامل آگهی را بچسبانید یا لینک معتبر بفرستید."
            ),
        }

    flow = pending.get("flow")
    if flow == "ai":
        base_cv = load_user_cv_text(telegram_user_id)
        if not base_cv:
            pending["step"] = "cv"
            _pending[int(telegram_user_id)] = pending
            return {
                "handled": True,
                "message": "رزومه پایه پیدا نشد. لطفاً فایل DOCX بفرستید.",
            }
        clear_pending(telegram_user_id)
        return {
            "handled": True,
            "message": "⏳ در حال ساخت رزومه با هوش مصنوعی… کمی صبر کنید.",
            "then_generate": {
                "telegram_user_id": telegram_user_id,
                "base_cv": base_cv,
                "title": job_title,
                "company": job_company,
                "job_url": job_url,
                "description": description,
                "mode": "ai",
            },
        }

    if flow == "apply_codex":
        from app.database import SessionLocal
        from app.worker.queue import JOB_CODEX_APPLY, enqueue_work, worker_status

        chat = int(pending.get("chat_id") or telegram_user_id)
        clear_pending(telegram_user_id)
        db = SessionLocal()
        try:
            ws = worker_status(db)
            work = enqueue_work(
                db,
                JOB_CODEX_APPLY,
                int(telegram_user_id),
                extra={
                    "title": job_title,
                    "company": job_company,
                    "job_url": job_url,
                    "description": description,
                    "chat_id": chat,
                    "telegram_user_id": int(telegram_user_id),
                },
                dedupe=True,
            )
        finally:
            db.close()

        online = bool(ws.get("worker_online"))
        msg = (
            "⏳ آگهی به صف <b>Codex / job-search-copilot</b> اضافه شد.\n"
            f"شناسه صف: <code>#{work.id}</code>\n\n"
        )
        if online:
            msg += "PC Worker آنلاین است — چند دقیقه صبر کنید تا نتیجه در همین چت بیاید."
        else:
            msg += (
                "⚠️ PC Worker آفلاین است.\n"
                "۱) تونل SSH را روشن کنید\n"
                "۲) <code>launch.bat</code> را اجرا کنید\n"
                "۳) اگر لازم است: <code>agent login</code> یا CURSOR_API_KEY"
            )
        return {"handled": True, "message": msg}

    if flow == "apply":
        base_cv = _resolve_base_cv(telegram_user_id, pending=pending)
        if len(base_cv) < 40:
            pending["step"] = "cv"
            _pending[int(telegram_user_id)] = pending
            return {
                "handled": True,
                "message": "رزومه پایه پیدا نشد. لطفاً فایل DOCX بفرستید.",
            }
        if pending.get("cv_text"):
            save_user_cv_text(telegram_user_id, base_cv)
        clear_pending(telegram_user_id)
        return {
            "handled": True,
            "message": (
                "⏳ در حال ارزیابی تناسب (job-search-copilot) و ساخت رزومه آماده درخواست…\n"
                "این مرحله ممکن است کمی طول بکشد."
            ),
            "then_generate": {
                "telegram_user_id": telegram_user_id,
                "base_cv": base_cv,
                "title": job_title,
                "company": job_company,
                "job_url": job_url,
                "description": description,
                "mode": "apply",
            },
        }

    if flow in ("ats", "combined"):
        base_cv = str(pending.get("cv_text") or "").strip() or (
            load_user_cv_text(telegram_user_id) or ""
        )
        if len(base_cv) < 40:
            pending["step"] = "cv"
            _pending[int(telegram_user_id)] = pending
            return {
                "handled": True,
                "message": "رزومه DOCX پیدا نشد. دوباره فایل را بفرستید.",
            }
        if flow == "combined":
            save_user_cv_text(telegram_user_id, base_cv)
        clear_pending(telegram_user_id)
        wait_msg = (
            "⏳ در حال ساخت رزومه AI + گزارش ATS… کمی صبر کنید."
            if flow == "combined"
            else "⏳ در حال بررسی ATS… کمی صبر کنید."
        )
        return {
            "handled": True,
            "message": wait_msg,
            "then_generate": {
                "telegram_user_id": telegram_user_id,
                "base_cv": base_cv,
                "title": job_title,
                "company": job_company,
                "job_url": job_url,
                "description": description,
                "mode": "combined" if flow == "combined" else "ats",
            },
        }

    clear_pending(telegram_user_id)
    return {"handled": True, "message": "عملیات نامشخص."}


async def run_generate_payload(payload: dict) -> dict:
    return await _run_generate(
        telegram_user_id=int(payload["telegram_user_id"]),
        base_cv=str(payload["base_cv"]),
        title=str(payload.get("title") or ""),
        company=str(payload.get("company") or ""),
        job_url=str(payload.get("job_url") or ""),
        description=str(payload.get("description") or ""),
        mode=str(payload.get("mode") or "ai"),
        prior_tips=payload.get("prior_tips") if isinstance(payload.get("prior_tips"), dict) else None,
        prior_scoring=payload.get("prior_scoring")
        if isinstance(payload.get("prior_scoring"), dict)
        else None,
        improve_pass=bool(payload.get("improve_pass")),
        previous_ats_score=int(payload.get("previous_ats_score") or 0),
    )


async def _resolve_job_input(raw: str) -> tuple[str, str, str, str]:
    """Return (title, company, url, description)."""
    job_id = parse_linkedin_job_id(raw)
    if job_id:
        from app.linkedin.search import fetch_job_posting

        meta = await fetch_job_posting(job_id)
        url = f"https://www.linkedin.com/jobs/view/{job_id}"
        title = (meta.get("title") or "").strip() or f"LinkedIn job {job_id}"
        company = (meta.get("company") or "").strip()
        description = (meta.get("description") or "").strip()
        location = (meta.get("location") or "").strip()
        if location and location.lower() not in description.lower()[:500]:
            description = f"Location: {location}\n\n{description}".strip()
        return title, company, url, description

    # Plain job description paste
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    title = lines[0][:120] if lines else "Target role"
    return title, "", "", raw


def _format_tips_message(tips: dict | None) -> str:
    if not isinstance(tips, dict):
        return ""
    tip_lines: list[str] = []
    summary = tips.get("summary")
    if summary:
        tip_lines.append(f"<b>راهنما</b>\n{_esc(summary)}")
    for tip in (tips.get("priority_tips") or [])[:6]:
        if isinstance(tip, dict):
            text = tip.get("tip") or tip.get("why") or ""
        else:
            text = tip
        if text:
            tip_lines.append(f"• {_esc(text)}")
    if not tip_lines:
        guide = tips.get("guide_text")
        if guide:
            tip_lines.append(_esc(str(guide)[:1500]))
    return "\n".join(tip_lines)


async def _run_generate(
    *,
    telegram_user_id: int,
    base_cv: str,
    title: str,
    company: str,
    job_url: str,
    description: str,
    mode: str,
    prior_tips: dict | None = None,
    prior_scoring: dict | None = None,
    improve_pass: bool = False,
    previous_ats_score: int = 0,
) -> dict:
    from app.ats.docx_export import export_resume_docx
    from app.ats.score import score_resume, score_resume_text
    from app.ats.tailor import tailor_resume_from_texts
    from app.ats.tips import build_improvement_guide
    from app.config import settings

    try:
        # ATS-only: score the user's uploaded resume — never build a new DOCX with AI
        if mode == "ats" and not improve_pass:
            scoring = score_resume_text(base_cv, description)
            tips = build_improvement_guide(scoring)
            band = scoring.get("band") or "—"
            total = int(scoring.get("total_score") or 0)
            matched = scoring.get("keyword_matched") or []
            missing = scoring.get("keyword_missing") or []
            _last_runs[int(telegram_user_id)] = {
                "base_cv": base_cv,
                "title": title,
                "company": company,
                "job_url": job_url,
                "description": description,
                "mode": "apply",  # improve will tailor with AI
                "tips": tips,
                "scoring": scoring,
            }
            caption = (
                "📊 <b>بررسی ATS آماده شد</b>\n\n"
                f"شغل: <b>{_esc(title)}</b>\n"
                f"امتیاز ATS: <b>{total}/100</b> · سطح: <b>{_esc(band)}</b>\n"
            )
            if matched:
                caption += f"کلیدواژه‌های منطبق: {_esc(', '.join(matched[:12]))}\n"
            if missing:
                caption += f"کلیدواژه‌های جاافتاده: {_esc(', '.join(missing[:12]))}\n"
            caption += (
                "\nبرای بازتولید رزومه بر اساس همین نقاط ضعف، "
                "دکمه <b>🔄 بهبود با نکات ATS</b> را بزنید."
            )
            return {
                "handled": True,
                "message": caption.strip(),
                "follow_up": _format_tips_message(tips) or None,
                "reply_markup": kb._markup(
                    [[kb._btn("🔄 بهبود با نکات ATS", "ut:improve")]]
                ),
            }

        fit_report = None
        fit = None
        if mode == "apply" and not improve_pass:
            from app.ats.fit_evaluate import evaluate_fit, format_fit_report

            fit = await evaluate_fit(
                base_cv=base_cv,
                title=title,
                description=description,
                company=company,
                job_url=job_url,
            )
            fit_report = format_fit_report(fit, title=title)
            if fit.get("recommendation") == "Skip" and fit.get("deal_breaker"):
                return {
                    "handled": True,
                    "message": fit_report
                    + "\n\n⏭ به‌خاطر deal-breaker رزومه ساخته نشد. اگر اصرار دارید دوباره /apply بزنید.",
                }

        resume = await tailor_resume_from_texts(
            base_cv=base_cv,
            title=title,
            description=description,
            company=company,
            job_url=job_url,
            prior_tips=prior_tips,
            prior_scoring=prior_scoring,
        )
        scoring = score_resume(resume, description, include_pdf=False)
        tips = build_improvement_guide(scoring)

        out_dir = USER_OUT_DIR / str(telegram_user_id)
        out_dir.mkdir(parents=True, exist_ok=True)
        safe = re.sub(r"[^A-Za-z0-9._-]+", "_", (resume.get("full_name") or "Resume"))[:40]
        suffix = "improved" if improve_pass else mode
        docx_path = out_dir / f"{safe}_{suffix}.docx"
        export_resume_docx(resume, docx_path)

        # Persist as BotApplication so FA translate / PDF buttons work
        import json as _json

        from app.config import ATS_DIR
        from app.database import BotApplication, SessionLocal
        from app.ats.pdf_export import export_resume_pdf

        db = SessionLocal()
        try:
            app = BotApplication(
                telegram_user_id=int(telegram_user_id),
                title=(title or "")[:512],
                company=(company or "")[:256],
                job_url=(job_url or "")[:1024],
                description=(description or "")[:60000],
                fit_score=int(fit.get("total_score") or 0) if fit else None,
                recommendation=(str(fit.get("recommendation") or "")[:64] if fit else None),
                ats_score=int(scoring.get("total_score") or 0),
                ats_scores_json=_json.dumps(scoring, ensure_ascii=False, default=str),
                resume_json=_json.dumps(resume, ensure_ascii=False, default=str)[:400000],
                status="ready",
            )
            db.add(app)
            db.flush()
            app_dir = Path(ATS_DIR) / "applications" / str(app.id)
            app_dir.mkdir(parents=True, exist_ok=True)
            app_docx = app_dir / "resume.docx"
            app_pdf = app_dir / "resume.pdf"
            export_resume_docx(resume, app_docx)
            try:
                export_resume_pdf(resume, app_pdf)
            except Exception:
                logger.exception("EN PDF failed in AI generate")
            if app_docx.is_file():
                app.resume_docx_path = str(app_docx)
                docx_path = app_docx
            if app_pdf.is_file():
                app.resume_pdf_path = str(app_pdf)
            db.commit()
            db.refresh(app)
            app_id = app.id
            markup = application_actions_markup(
                app,
                improve_btn=kb._btn("🔄 بهبود با نکات ATS", "ut:improve"),
            )
        finally:
            db.close()

        band = scoring.get("band") or "—"
        total = int(scoring.get("total_score") or 0)
        matched = scoring.get("keyword_matched") or []
        missing = scoring.get("keyword_missing") or []
        diff = _esc(resume.get("diff_summary") or "")
        provider = settings.proposal_provider
        model = settings.proposal_model()

        _last_runs[int(telegram_user_id)] = {
            "base_cv": base_cv,
            "title": title,
            "company": company,
            "job_url": job_url,
            "description": description,
            "mode": mode if mode != "ats" else "apply",
            "tips": tips,
            "scoring": scoring,
            "application_id": app_id,
        }

        if improve_pass:
            header = "🔄 <b>رزومه بهبودیافته با نکات ATS</b>"
        else:
            header = {
                "ai": "📄 <b>رزومه با AI آماده شد</b>",
                "combined": "📄📊 <b>رزومه AI + گزارش ATS آماده شد</b>",
                "apply": "✅ <b>بسته درخواست آماده شد</b>",
            }.get(mode, "📄 <b>رزومه آماده شد</b>")
        caption = f"{header}\n\nشغل: <b>{_esc(title)}</b>\n"
        if fit:
            caption += (
                f"تناسب: <b>{int(fit.get('total_score') or 0)}/100</b> · "
                f"توصیه: <b>{_esc(fit.get('recommendation'))}</b>\n"
            )
        caption += f"امتیاز ATS: <b>{total}/100</b> · سطح: <b>{_esc(band)}</b>\n"
        if improve_pass and previous_ats_score:
            delta = total - previous_ats_score
            arrow = "▲" if delta > 0 else ("▼" if delta < 0 else "•")
            caption += (
                f"نسبت به قبل: <b>{previous_ats_score}</b> → <b>{total}</b> "
                f"({arrow}{abs(delta)})\n"
            )
        if matched:
            caption += f"کلیدواژه‌های منطبق: {_esc(', '.join(matched[:12]))}\n"
        if missing:
            caption += f"کلیدواژه‌های جاافتاده: {_esc(', '.join(missing[:12]))}\n"
        caption += f"AI: <code>{_esc(provider)}</code> / <code>{_esc(model)}</code>\n"
        if diff:
            caption += f"\n{diff}"
        caption += (
            "\n\nدکمه‌ها: بهبود ATS · ترجمه فارسی (DOCX RTL) · PDF"
        )
        if mode == "apply" or improve_pass:
            caption += (
                "\n📦 فایل DOCX آماده است. قبل از ارسال واقعی مرور کنید."
            )

        follow_parts = []
        if fit_report:
            follow_parts.append(fit_report)
        tips_msg = _format_tips_message(tips)
        if tips_msg:
            follow_parts.append(tips_msg)

        return {
            "handled": True,
            "message": caption,
            "follow_up": "\n\n——\n\n".join(follow_parts) if follow_parts else None,
            "send_document": str(docx_path),
            "document_caption": (
                f"Fit {int(fit.get('total_score') or 0)}/100 · {fit.get('recommendation')} · "
                f"ATS {total}/100"
                if fit
                else f"امتیاز ATS: {total}/100 · {band}"
            ),
            "reply_markup": markup,
            "skip_improve_hint": True,
        }
    except Exception as exc:
        logger.exception("User tool generate failed (%s)", mode)
        fail = (
            "❌ بررسی ATS ممکن نشد"
            if mode == "ats" and not improve_pass
            else "❌ ساخت رزومه / ارزیابی ممکن نشد"
        )
        return {
            "handled": True,
            "message": f"{fail}: {_esc(exc)}",
        }
